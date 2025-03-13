import logging
import re
from docx.enum.text import WD_COLOR_INDEX
import copy

logger = logging.getLogger(__name__)


def rebuild_document(doc, elements, processed_elements):
    """
    Rebuild the document with processed elements
    
    Args:
        doc: The Word document to rebuild
        elements: The original document elements
        processed_elements: The processed text for each element
        
    Returns:
        str: The full simplified text of the document
    """
    logger.info("🔄 Rebuilding document with simplified text...")
    
    # Separate media elements from text elements
    text_elements = [e for e in elements if e["type"] != "media"]
    media_elements = [e for e in elements if e["type"] == "media"]
    
    logger.info(f"Found {len(text_elements)} text elements and {len(media_elements)} media elements")
    
    # Create a set of paragraph indexes that contain media elements
    media_paragraph_indexes = set()
    for media_element in media_elements:
        if "index" in media_element:
            media_paragraph_indexes.add(media_element["index"])
    
    logger.info(f"Found {len(media_paragraph_indexes)} paragraphs containing media elements")
    
    # Create a mapping of element indexes to original paragraphs
    paragraph_map = {}
    for i, para in enumerate(doc.paragraphs):
        paragraph_map[i] = para
    
    # Only clear paragraphs that don't contain media elements
    for i, paragraph in enumerate(doc.paragraphs):
        if i not in media_paragraph_indexes:
            paragraph.clear()
    
    # Rebuild each paragraph with its processed text
    simplified_full_text = ""
    
    for element in text_elements:
        element_id = element["id"]
        original_index = element["index"]
        
        # Skip paragraphs that contain media elements
        if original_index in media_paragraph_indexes:
            simplified_full_text += element["text"] + "\n"
            continue
        
        # Get the processed text for this element
        processed_text = processed_elements.get(element_id, element["text"])
        
        # Ensure processed_text is not None before concatenating
        if processed_text is not None:
            simplified_full_text += processed_text + "\n"
        else:
            simplified_full_text += element["text"] + "\n"
            processed_text = element["text"]  # Use original text if processed text is None
        
        # Get the original paragraph
        paragraph = paragraph_map.get(original_index)
        if not paragraph:
            continue
        
        # Apply text and restore formatting
        restore_paragraph_formatting(paragraph, element, processed_text)
    
    logger.info(f"✅ Document rebuilt with {len(text_elements)} text elements and {len(media_elements)} preserved media elements")
    return simplified_full_text

def restore_paragraph_formatting(paragraph, element, processed_text):
    """
    Restore formatting to a paragraph based on the original element
    
    Args:
        paragraph: The Word paragraph to format
        element: The original document element
        processed_text: The processed text for this element
    """
    # Set paragraph text
    if not processed_text:
        return
        
    # Determine if the element had any highlighted text
    highlighted_phrases = element["highlighted_phrases"]
    
    # If it's a heading, preserve full formatting
    if element["type"].startswith("heading_") or element["type"] in ["toc_entry", "caption"]:
        # Set paragraph alignment
        if element["format"]["alignment"] is not None:
            paragraph.alignment = element["format"]["alignment"]
            
        # Create a single run with the paragraph's text
        run = paragraph.add_run(processed_text)
        
        # Apply formatting from original runs
        if element["runs"] and len(element["runs"]) > 0:
            for run_data in element["runs"]:
                if run_data["bold"]:
                    run.bold = True
                    break
                
        return
        
    # For list items, keep the original list marker formatting
    if element["type"] == "list_item":
        # Extract list marker if present
        list_marker = ""
        text = processed_text
        
        # Check for bullet or number markers
        match = re.match(r'^(\s*[-•*]|\s*\d+\.|\s*[a-zA-Z]\.)\s+', processed_text)
        if match:
            list_marker = match.group(1)
            text = processed_text[len(match.group(0)):]
            
            # Add the list marker with original formatting
            run = paragraph.add_run(list_marker + " ")
            paragraph.alignment = element["format"]["alignment"]
            processed_text = text  # Continue with the rest of the text
    
    # For regular paragraphs or the remaining text in list items
    # Try to map original formatting to the simplified text
    
    # Check if we can do word-by-word formatting mapping
    if len(element["runs"]) > 1 and not all(r["bold"] == element["runs"][0]["bold"] and 
                                          r["italic"] == element["runs"][0]["italic"] and
                                          r["underline"] == element["runs"][0]["underline"] 
                                          for r in element["runs"]):
        # Complex case - mixed formatting within the paragraph
        apply_complex_formatting(paragraph, element, processed_text)
    else:
        # Simple case - consistent formatting or no special formatting
        if not highlighted_phrases:
            # Set paragraph alignment
            if element["format"]["alignment"] is not None:
                paragraph.alignment = element["format"]["alignment"]
                
            # Create a single run with the paragraph's text
            run = paragraph.add_run(processed_text)
            
            # Apply consistent formatting if present
            if element["runs"] and all(r["bold"] for r in element["runs"]):
                run.bold = True
            if element["runs"] and all(r["italic"] for r in element["runs"]):
                run.italic = True
            if element["runs"] and all(r["underline"] for r in element["runs"]):
                run.underline = True
        else:
            # Handle text with highlights
            text_to_write = processed_text
            
            # Mark highlighted phrases for later processing
            for phrase in highlighted_phrases:
                if phrase in processed_text:
                    text_to_write = text_to_write.replace(phrase, f"[highlight]{phrase}[/highlight]")
            
            # Process text with highlights
            parts = text_to_write.split("[highlight]")
            
            # Set paragraph alignment
            if element["format"]["alignment"] is not None:
                paragraph.alignment = element["format"]["alignment"]
                
            for part in parts:
                if "[/highlight]" in part:
                    highlighted, rest = part.split("[/highlight]")
                    run = paragraph.add_run(highlighted)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Apply consistent formatting if present
                    if element["runs"] and all(r["bold"] for r in element["runs"]):
                        run.bold = True
                    if element["runs"] and all(r["italic"] for r in element["runs"]):
                        run.italic = True
                    
                    if rest:
                        run = paragraph.add_run(rest)
                        # Apply consistent formatting if present
                        if element["runs"] and all(r["bold"] for r in element["runs"]):
                            run.bold = True
                        if element["runs"] and all(r["italic"] for r in element["runs"]):
                            run.italic = True
                else:
                    run = paragraph.add_run(part)
                    # Apply consistent formatting if present
                    if element["runs"] and all(r["bold"] for r in element["runs"]):
                        run.bold = True
                    if element["runs"] and all(r["italic"] for r in element["runs"]):
                        run.italic = True

def apply_complex_formatting(paragraph, element, processed_text):
    """
    Apply mixed formatting from the original runs to the processed text
    
    Args:
        paragraph: The Word paragraph to format
        element: The original document element
        processed_text: The processed text for this element
    """
    # Set paragraph alignment
    if element["format"]["alignment"] is not None:
        paragraph.alignment = element["format"]["alignment"]
        
    # Build a mapping of formatting styles in the original text
    original_text = element["text"]
    
    # If the simplified text is too different from the original, fall back to simpler formatting
    if len(processed_text) < 0.5 * len(original_text) or len(processed_text) > 1.5 * len(original_text):
        # Create a single run with default formatting
        run = paragraph.add_run(processed_text)
        return
    
    # Try to identify key words and their formatting from the original
    formatted_words = {}
    
    for run_data in element["runs"]:
        text = run_data["text"]
        if not text:
            continue
            
        # Clean and normalize the text for better matching
        clean_text = text.strip().lower()
        if not clean_text:
            continue
            
        # Store formatting for each word
        for word in re.findall(r'\b\w+\b', clean_text):
            if len(word) > 3:  # Only track substantial words
                formatted_words[word] = {
                    "bold": run_data["bold"],
                    "italic": run_data["italic"],
                    "underline": run_data["underline"],
                    "highlight": run_data["highlight"]
                }
    
    # Now try to apply this formatting to words in the processed text
    
    # Split the processed text into words and non-word segments
    segments = re.findall(r'(\b\w+\b|\W+)', processed_text)
    
    for segment in segments:
        # Check if this is a word we should format
        clean_segment = segment.strip().lower()
        
        if clean_segment in formatted_words and len(clean_segment) > 3:
            # This word had special formatting in the original
            run = paragraph.add_run(segment)
            format_info = formatted_words[clean_segment]
            
            if format_info["bold"]:
                run.bold = True
            if format_info["italic"]:
                run.italic = True
            if format_info["underline"]:
                run.underline = True
            if format_info["highlight"]:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            # No special formatting for this segment
            run = paragraph.add_run(segment) 
